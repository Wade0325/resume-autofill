"""
docx 錨點抽取器 (Anchor Extractor)
------------------------------------------------
把任何一份 Word 履歷表拆解成一串「可以填東西的位置」＋「它旁邊的標籤文字」。

這一步刻意「完全不用 AI」，純結構解析。原因：
  * Word 的版面資訊（表格座標、儲存格、內容控制項）是精確的，交給程式最可靠
  * AI 只需要做一件事：判斷「這個標籤是什麼欄位」，那是純語意題，小模型就夠

支援的四種錨點型態：
  sdt        內容控制項 (Content Control / w:sdt) —— 最好填，有 tag/alias
  formfield  舊式表單欄位 (FORMTEXT) —— 常見於老公司的範本
  cell       表格空白儲存格 —— 台灣公司履歷表最常見（左邊或上面是標籤）
  inline     段落內的填空 —— 「姓名：______」「電話：」
  checkbox   勾選題 —— 「性別 □男 □女」
"""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass, asdict, field
from typing import Any, Dict, Iterator, List, Optional

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# 常見的「這裡是空白」標記
PLACEHOLDER_RE = re.compile(r"^[\s\u3000_＿…．\.\-—–]*$")
BLANK_RUN_RE = re.compile(r"[_＿]{2,}|[\.．]{4,}")
CHECKBOX_CHARS = "□☐▢◻"   # 刻意不含 ○〇◯：中文常用來遮蔽名稱（○○公司）
CHECKED_CHARS = "■☑▣◼"
# 「標籤：值」形式，值可為空
INLINE_LABEL_RE = re.compile(
    r"(?P<label>[\u4e00-\u9fffA-Za-z][\u4e00-\u9fffA-Za-z0-9 \u3000（）()／/、\-]{0,14})"
    r"\s*[:：]\s*(?P<value>[^\n]*)$"
)


@dataclass
class Anchor:
    """一個可填寫的位置。"""
    id: str                       # 穩定識別碼，例如 "tbl0.r2.c1"
    kind: str                     # sdt | formfield | cell | inline | checkbox
    label: str                    # 推測出來的標籤文字（給比對用）
    context: str                  # 周邊文字（給 LLM 判斷用）
    loc: Dict[str, Any] = field(default_factory=dict)   # 精確座標，給 writer 用
    options: List[str] = field(default_factory=list)    # checkbox 選項
    existing: str = ""            # 目前已有的內容（若非空表示已填過）

    def to_dict(self):
        return asdict(self)


# --------------------------------------------------------------------------
# 文件走訪工具
# --------------------------------------------------------------------------
def iter_block_items(parent) -> Iterator[Any]:
    """依照文件真實順序走訪段落與表格（python-docx 官方 recipe 的變體）。"""
    if isinstance(parent, _Doc):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"不支援的容器：{type(parent)}")
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def is_blank(text: str) -> bool:
    return bool(PLACEHOLDER_RE.match(text or ""))


# --------------------------------------------------------------------------
# 1) 內容控制項 / 舊式表單欄位
# --------------------------------------------------------------------------
def extract_content_controls(doc) -> List[Anchor]:
    anchors: List[Anchor] = []
    body = doc.element.body
    for i, sdt in enumerate(body.iter(qn("w:sdt"))):
        pr = sdt.find(qn("w:sdtPr"))
        tag = alias = ""
        if pr is not None:
            t = pr.find(qn("w:tag"))
            a = pr.find(qn("w:alias"))
            tag = t.get(qn("w:val")) if t is not None else ""
            alias = a.get(qn("w:val")) if a is not None else ""
        content = sdt.find(qn("w:sdtContent"))
        cur = "".join(n.text or "" for n in content.iter(qn("w:t"))) if content is not None else ""
        anchors.append(Anchor(
            id=f"sdt{i}",
            kind="sdt",
            label=(alias or tag or "").strip(),
            context=f"內容控制項 alias={alias!r} tag={tag!r} 目前值={cur!r}",
            loc={"sdt_index": i},
            existing=cur,
        ))
    return anchors


def extract_form_fields(doc) -> List[Anchor]:
    """舊式 FORMTEXT 欄位：<w:fldChar w:fldCharType="begin"><w:ffData><w:name .../>"""
    anchors: List[Anchor] = []
    body = doc.element.body
    idx = 0
    for ff in body.iter(qn("w:ffData")):
        name_el = ff.find(qn("w:name"))
        name = name_el.get(qn("w:val")) if name_el is not None else ""
        default = ""
        text_input = ff.find(qn("w:textInput"))
        if text_input is not None:
            d = text_input.find(qn("w:default"))
            if d is not None:
                default = d.get(qn("w:val")) or ""
        anchors.append(Anchor(
            id=f"ff{idx}",
            kind="formfield",
            label=(name or "").strip(),
            context=f"舊式表單欄位 name={name!r} 預設值={default!r}",
            loc={"ff_index": idx, "name": name},
            existing=default,
        ))
        idx += 1
    return anchors


# --------------------------------------------------------------------------
# 2) 表格
# --------------------------------------------------------------------------
def _grid(table: Table) -> List[List[Optional[_Cell]]]:
    """回傳去重後的儲存格網格；合併儲存格只在左上角保留一次，其餘為 None。"""
    seen = set()
    grid: List[List[Optional[_Cell]]] = []
    for row in table.rows:
        line: List[Optional[_Cell]] = []
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:
                line.append(None)
            else:
                seen.add(key)
                line.append(cell)
        grid.append(line)
    return grid


def _label_for_cell(grid, r: int, c: int) -> str:
    """空白格的標籤：先看左邊，再看上面（台灣履歷表兩種都很常見）。"""
    for cc in range(c - 1, -1, -1):
        cell = grid[r][cc]
        if cell is not None:
            t = cell_text(cell)
            if t and not is_blank(t):
                return t.replace("\n", " ")
    for rr in range(r - 1, -1, -1):
        if c < len(grid[rr]):
            cell = grid[rr][c]
            if cell is not None:
                t = cell_text(cell)
                if t and not is_blank(t):
                    return t.replace("\n", " ")
    return ""


def _row_context(grid, r: int) -> str:
    texts = []
    for cell in grid[r]:
        if cell is not None:
            t = cell_text(cell).replace("\n", " ")
            texts.append(t if t else "␣")
    return " | ".join(texts)[:300]


def table_preceding_labels(doc) -> Dict[int, str]:
    """回傳 {表格序號: 緊接在它前面的那段非空文字}。
    用來處理「四、自傳」後面接一個空白單格表格這種常見版型。"""
    labels: Dict[int, str] = {}
    last_text = ""
    ti = 0
    for block in iter_block_items(doc):
        if isinstance(block, Table):
            labels[ti] = last_text
            ti += 1
        else:
            t = block.text.strip()
            if t:
                last_text = t
    return labels


def extract_tables(doc, prefix="tbl") -> List[Anchor]:
    anchors: List[Anchor] = []
    pre_labels = table_preceding_labels(doc)
    for ti, table in enumerate(doc.tables):
        grid = _grid(table)
        for r, line in enumerate(grid):
            for c, cell in enumerate(line):
                if cell is None:
                    continue
                text = cell_text(cell)
                aid = f"{prefix}{ti}.r{r}.c{c}"
                loc = {"table": ti, "row": r, "col": c}

                # (a) 勾選題
                if any(ch in text for ch in CHECKBOX_CHARS):
                    opts = _parse_checkbox_options(text)
                    label = _strip_options(text) or _label_for_cell(grid, r, c)
                    anchors.append(Anchor(
                        id=aid, kind="checkbox", label=label,
                        context=_row_context(grid, r), loc=loc,
                        options=opts, existing=text))
                    continue

                # (b) 空白格 → 左/上找標籤
                if is_blank(text):
                    label = _label_for_cell(grid, r, c)
                    if not label:
                        # 網格內找不到標籤 → 退而求其次用表格前面那段標題
                        label = _clean_heading(pre_labels.get(ti, ""))
                    if not label:
                        continue
                    anchors.append(Anchor(
                        id=aid, kind="cell", label=label,
                        context=_row_context(grid, r), loc=loc, existing=""))
                    continue

                # (c) 同格內「標籤：」後面留白
                m = INLINE_LABEL_RE.search(text.split("\n")[-1])
                if m and is_blank(m.group("value")):
                    anchors.append(Anchor(
                        id=aid + ".inline", kind="inline",
                        label=m.group("label").strip(),
                        context=_row_context(grid, r),
                        loc={**loc, "in_cell": True}, existing=""))
    return anchors


def _parse_checkbox_options(text: str) -> List[str]:
    """從『性別 □男 □女』抓出 ['男','女']。"""
    opts = []
    pattern = re.compile(f"[{CHECKBOX_CHARS}{CHECKED_CHARS}]\\s*([^\\s{CHECKBOX_CHARS}{CHECKED_CHARS}]{{1,10}})")
    for m in pattern.finditer(text):
        v = m.group(1).strip(" \u3000、,，/:：")
        if v:
            opts.append(v)
    return opts


def _strip_options(text: str) -> str:
    return re.sub(f"[{CHECKBOX_CHARS}{CHECKED_CHARS}][^\\s]*", "", text).strip()


# --------------------------------------------------------------------------
# 3) 段落填空
# --------------------------------------------------------------------------
def extract_paragraphs(doc) -> List[Anchor]:
    anchors: List[Anchor] = []
    for pi, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue

        if any(ch in text for ch in CHECKBOX_CHARS):
            opts = _parse_checkbox_options(text)
            if opts:
                anchors.append(Anchor(
                    id=f"p{pi}.chk", kind="checkbox",
                    label=_strip_options(text)[:30],
                    context=text[:200], loc={"para": pi},
                    options=opts, existing=text))
                continue

        # 一行可能有好幾個填空：「緊急聯絡人：___ 關係：___ 電話：___」
        blanks = list(BLANK_RUN_RE.finditer(text))
        if blanks:
            prev_end = 0
            for bi, m in enumerate(blanks):
                label = _label_before(text, prev_end, m.start())
                prev_end = m.end()
                if not label:
                    continue
                anchors.append(Anchor(
                    id=f"p{pi}.b{bi}", kind="inline", label=label,
                    context=text[:200],
                    loc={"para": pi, "blank_index": bi}))
            continue

        # 「可到職日：」冒號結尾、後面直接留白
        m = re.search(r"(?P<label>[^\s\u3000:：]{1,14})[:：][ \u3000]*$", text)
        if m:
            anchors.append(Anchor(
                id=f"p{pi}.tail", kind="inline",
                label=m.group("label").strip(),
                context=text[:200], loc={"para": pi, "tail": True}))
    return anchors


def _label_before(text: str, start: int, end: int) -> str:
    """取填空左邊那一小段當標籤：『… 關係：』→『關係』。"""
    seg = text[start:end].strip()
    seg = seg.rstrip(" \u3000:：")
    seg = re.split(r"[\s\u3000,，;；]", seg)[-1] if seg else ""
    return seg[-20:]


def _clean_heading(text: str) -> str:
    """把「四、自傳」清成「自傳」。"""
    return re.sub(r"^[\s\u3000]*[一二三四五六七八九十０-９0-9]+[、.．)）]?[\s\u3000]*", "",
                  (text or "").strip())


# --------------------------------------------------------------------------
# 對外主函式
# --------------------------------------------------------------------------
def extract(path: str) -> Dict[str, Any]:
    doc = Document(path)
    anchors: List[Anchor] = []
    anchors += extract_content_controls(doc)
    anchors += extract_form_fields(doc)
    anchors += extract_tables(doc)
    anchors += extract_paragraphs(doc)

    # 去重（同一個 id 只留一個）
    uniq: Dict[str, Anchor] = {}
    for a in anchors:
        uniq.setdefault(a.id, a)
    anchors = list(uniq.values())

    return {
        "path": path,
        "fingerprint": fingerprint(anchors),
        "anchors": [a.to_dict() for a in anchors],
    }


def fingerprint(anchors: List[Anchor]) -> str:
    """
    範本指紋：同一家公司的同一份表格 → 同一個指紋 → 直接沿用上次的對映結果，
    第二次以後完全不用叫 LLM，秒填。
    """
    payload = json.dumps(
        [[a.kind, a.label, a.id] for a in sorted(anchors, key=lambda x: x.id)],
        ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]
