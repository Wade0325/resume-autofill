"""產生一份「標籤刻意寫得很怪」的履歷表，用來測 LLM 那一層。

make_sample.py 的標準標籤（姓名、行動電話）就算模型判錯，
確定性標籤對齊也會拉回來。這一份的標籤在別名表裡都找不到，
對映結果完全取決於模型本身的判斷。
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(path="tricky_resume_form.docx"):
    doc = Document()
    doc.styles["Normal"].font.name = "DFKai-SB"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_paragraph("△△國際企業　人才招募申請書")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(16)

    t = doc.add_table(rows=6, cols=4)
    t.style = "Table Grid"
    layout = [
        ["應徵者大名", "", "生理性別", "□男　□女"],
        ["出生之年月日", "", "日間可聯絡之電話", ""],
        ["電子郵遞地址", "", "服役情況", "□役畢　□免役　□未役"],
        ["現居住所", "", "", ""],
        ["身分證上之住址", "", "", ""],
        ["欲應徵之職位", "", "期望之薪資待遇", ""],
    ]
    for r, row in enumerate(layout):
        for c, val in enumerate(row):
            t.cell(r, c).text = val
    t.cell(3, 1).merge(t.cell(3, 3))
    t.cell(4, 1).merge(t.cell(4, 3))

    doc.add_paragraph()
    doc.add_paragraph("壹、求學歷程").runs[0].bold = True
    edu = doc.add_table(rows=3, cols=4)
    edu.style = "Table Grid"
    for c, head in enumerate(["肄業或畢業之學校", "所屬院系所", "取得之學位", "起訖年月"]):
        edu.cell(0, c).text = head

    doc.add_paragraph()
    doc.add_paragraph("貳、服務資歷").runs[0].bold = True
    exp = doc.add_table(rows=3, cols=4)
    exp.style = "Table Grid"
    for c, head in enumerate(["服務機構全銜", "擔任之職務", "在職起訖", "去職緣由"]):
        exp.cell(0, c).text = head

    doc.add_paragraph()
    doc.add_paragraph("參、其他事項").runs[0].bold = True
    doc.add_paragraph("外語能力：____________________________")
    doc.add_paragraph("持有之專業證書：____________________________")
    doc.add_paragraph("急難時聯絡之人：____________  與本人之關係：________  其聯絡電話：____________")
    doc.add_paragraph("何時可以來上班：")

    doc.add_paragraph()
    doc.add_paragraph("肆、個人簡述").runs[0].bold = True
    bio = doc.add_table(rows=1, cols=1)
    bio.style = "Table Grid"
    bio.cell(0, 0).text = ""

    doc.add_paragraph()
    doc.add_paragraph("※以下欄位請勿填寫")
    hr = doc.add_table(rows=1, cols=2)
    hr.style = "Table Grid"
    hr.cell(0, 0).text = "初審意見"
    hr.cell(0, 1).text = ""

    doc.save(path)
    print(f"✓ 已產生 {path}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "tricky_resume_form.docx")
