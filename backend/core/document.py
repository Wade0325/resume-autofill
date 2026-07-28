"""docx 的讀寫基礎：把文件攤成文字，並列出可以填字的位置。

這裡刻意不判斷任何一格「是什麼欄位」——那是模型的工作。
本模組只回答兩件機械性的問題：文件寫了什麼、哪些位置可以寫字、寫在哪個座標。
"""
from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass, field
from typing import Any, Dict, Iterator, List, Optional, Set, Tuple

from docx import Document
from docx.document import Document as _Doc
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

PLACEHOLDER_RE = re.compile(r"^[\s　_＿…．\.\-—–]*$")
BLANK_RUN_RE = re.compile(r"[_＿]{2,}|[\.．]{4,}")
CHECKBOX_CHARS = "□☐▢◻"   # 不含 ○〇◯：中文常用來遮蔽名稱（○○公司）
CHECKED_CHARS = "■☑▣◼"
TRAILING_COLON_RE = re.compile(r"[:：][ 　]*$")


@dataclass
class Slot:
    """一個可以填字的位置。"""
    id: str
    kind: str                                            # sdt | formfield | cell | inline | checkbox
    loc: Dict[str, Any] = field(default_factory=dict)    # writer 用來定位
    options: List[str] = field(default_factory=list)     # 勾選題的選項
    existing: str = ""                                   # 目前的內容，非空代表會被覆蓋

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


# --------------------------------------------------------------------------
# 走訪工具
# --------------------------------------------------------------------------
def iter_block_items(parent) -> Iterator[Any]:
    body = parent.element.body if isinstance(parent, _Doc) else parent._tc
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def is_blank(text: str) -> bool:
    return bool(PLACEHOLDER_RE.match(text or ""))


def squash(text: str) -> str:
    """比對用：抹掉空白差異。"""
    return re.sub(r"[\s　]+", "", text or "")


def checkbox_options(text: str) -> List[str]:
    """從「□男　□女」取出選項文字。"""
    marks = CHECKBOX_CHARS + CHECKED_CHARS
    parts = re.split(f"[{marks}]", text)[1:]
    out = []
    for part in parts:
        name = re.split(r"[\s　,，、/／]+", part.strip())[0].strip(" :：()（）")
        if name:
            out.append(name)
    return out


def checked_option(text: str) -> str:
    """取出已經被勾選的那一項。"""
    m = re.search(f"[{CHECKED_CHARS}][ 　]*([^ 　{CHECKBOX_CHARS}{CHECKED_CHARS}]+)", text)
    return m.group(1).strip() if m else ""


# --------------------------------------------------------------------------
# 主要出口
# --------------------------------------------------------------------------
def load(path: str, overwritable: Optional[Set[str]] = None) -> Tuple[str, List[Slot]]:
    """回傳（帶位置標記的全文, 位置清單）。

    空白的位置一律可填。已經有字的格子只有在 overwritable 裡才算可填——
    表格印好的欄位名稱與使用者填的值長得一樣（都是非空儲存格），
    差別只在後者是這個人的資料。那份清單由 reader 讀出來，所以判斷依據
    仍然是模型，不是規則。
    """
    doc = Document(path)
    overwritable = {squash(v) for v in (overwritable or set())}
    slots: List[Slot] = []
    lines: List[str] = []

    slots += _content_controls(doc)
    slots += _form_fields(doc)

    # 段落序號必須是 doc.paragraphs 的索引，writer 靠它定位
    table_index = 0
    para_index = 0
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            line = _paragraph_line(block, para_index, slots)
            if line:
                lines.append(line)
            para_index += 1
        elif isinstance(block, Table):
            lines.append(f"[表格{table_index}]")
            lines += _table_lines(block, table_index, slots, overwritable)
            table_index += 1

    return "\n".join(lines), slots


def text_only(path: str) -> str:
    """只要全文，不標記位置（匯入用）。"""
    return load(path)[0]


def slot_headers(path: str, slots: List[Slot]) -> Dict[str, Dict[str, str]]:
    """每個表格位置的「列首」與「欄首」——同列往左、同欄往上第一格有字的內容。

    這是機械抽取，不經過模型，所以可靠。給第二輪審核用：
    「列首＝高中/專科、欄首＝學校名稱」比整份攤平全文精準得多。
    """
    doc = Document(path)
    tables = [b for b in iter_block_items(doc) if isinstance(b, Table)]

    by_table: Dict[int, List[Slot]] = {}
    for s in slots:
        if "table" in s.loc:
            by_table.setdefault(s.loc["table"], []).append(s)

    out: Dict[str, Dict[str, str]] = {}
    for ti, tslots in by_table.items():
        if ti >= len(tables):
            continue
        texts = [[cell_text(c) if c is not None else "" for c in row]
                 for row in _grid(tables[ti])]
        for s in tslots:
            r, c = s.loc["row"], s.loc["col"]
            if r >= len(texts):
                continue
            row_hdr = next(
                (t for t in reversed(texts[r][:c]) if t.strip() and not is_blank(t)), "")
            col_hdr = ""
            for rr in range(r - 1, -1, -1):
                if c < len(texts[rr]) and texts[rr][c].strip() and not is_blank(texts[rr][c]):
                    col_hdr = texts[rr][c]
                    break
            out[s.id] = {"row": row_hdr.replace("\n", " ")[:20],
                         "col": col_hdr.replace("\n", " ")[:20]}
    return out


def fingerprint(slots: List[Slot]) -> str:
    """同一份表格 → 同一個指紋 → 直接沿用上次的對映，不必再問模型。"""
    payload = json.dumps(sorted((s.kind, s.id) for s in slots), ensure_ascii=False)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]


# --------------------------------------------------------------------------
# 各種位置
# --------------------------------------------------------------------------
def _content_controls(doc) -> List[Slot]:
    out = []
    for i, sdt in enumerate(doc.element.body.iter(qn("w:sdt"))):
        texts = [t.text or "" for t in sdt.iter(qn("w:t"))]
        out.append(Slot(id=f"sdt{i}", kind="sdt", loc={"sdt_index": i},
                        existing="".join(texts).strip()))
    return out


def _form_fields(doc) -> List[Slot]:
    out = []
    for i, ff in enumerate(doc.element.body.iter(qn("w:ffData"))):
        name_el = ff.find(qn("w:name"))
        text_el = ff.find(qn("w:textInput"))
        default = ""
        if text_el is not None:
            default_el = text_el.find(qn("w:default"))
            if default_el is not None:
                default = default_el.get(qn("w:val"), "")
        out.append(Slot(id=f"ff{i}", kind="formfield",
                        loc={"ff_index": i,
                             "name": name_el.get(qn("w:val"), "") if name_el is not None else ""},
                        existing=default))
    return out


def _paragraph_line(para: Paragraph, index: int, slots: List[Slot]) -> str:
    text = para.text.strip()
    if not text:
        return ""

    if any(ch in text for ch in CHECKBOX_CHARS + CHECKED_CHARS):
        options = checkbox_options(text)
        if options:
            sid = f"p{index}.chk"
            slots.append(Slot(id=sid, kind="checkbox", loc={"para": index},
                              options=options, existing=text))
            return f"{{{{{sid}}}}} {text}"

    blanks = list(BLANK_RUN_RE.finditer(text))
    if blanks:
        out, cursor = [], 0
        for bi, m in enumerate(blanks):
            sid = f"p{index}.b{bi}"
            slots.append(Slot(id=sid, kind="inline",
                              loc={"para": index, "blank_index": bi}))
            out.append(text[cursor:m.start()] + f"{{{{{sid}}}}}")
            cursor = m.end()
        return "".join(out) + text[cursor:]

    if TRAILING_COLON_RE.search(text):
        sid = f"p{index}.tail"
        slots.append(Slot(id=sid, kind="inline", loc={"para": index, "tail": True}))
        return f"{text}{{{{{sid}}}}}"

    return text


def _table_lines(table: Table, table_index: int, slots: List[Slot],
                 overwritable: Set[str]) -> List[str]:
    out = []
    for r, row in enumerate(_grid(table)):
        rendered, seen = [], set()
        for c, cell in enumerate(row):
            if cell is None:
                continue
            # 合併儲存格會在多個座標回傳同一個物件
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            rendered.append(_cell_render(cell, table_index, r, c, slots, overwritable))
        if any(x.strip() and x.strip() != "␣" for x in rendered):
            out.append(" | ".join(rendered))
    return out


def _cell_render(cell: _Cell, table_index: int, r: int, c: int,
                 slots: List[Slot], overwritable: Set[str]) -> str:
    text = cell_text(cell)
    sid = f"tbl{table_index}.r{r}.c{c}"
    loc = {"table": table_index, "row": r, "col": c}

    if any(ch in text for ch in CHECKBOX_CHARS + CHECKED_CHARS):
        options = checkbox_options(text)
        if options:
            slots.append(Slot(id=sid, kind="checkbox", loc=loc,
                              options=options, existing=text))
            return f"{{{{{sid}}}}} {text}".replace("\n", " ")

    if is_blank(text):
        slots.append(Slot(id=sid, kind="cell", loc=loc))
        return f"{{{{{sid}}}}}"

    if squash(text) in overwritable:
        slots.append(Slot(id=sid, kind="cell", loc=loc, existing=text))
        return f"{{{{{sid}}}}}{text}".replace("\n", " ")

    return text.replace("\n", " ")


def _grid(table: Table) -> List[List[_Cell]]:
    """把合併儲存格攤平成矩形網格，同一個 cell 會在它涵蓋的每個座標出現。"""
    rows: List[List[_Cell]] = []
    for row in table.rows:
        cells: List[_Cell] = []
        for tc in row._tr.tc_lst:
            cell = _Cell(tc, table)
            span = tc.tcPr.find(qn("w:gridSpan")) if tc.tcPr is not None else None
            width = int(span.get(qn("w:val"))) if span is not None else 1
            cells.extend([cell] * width)
        rows.append(cells)
    return rows
