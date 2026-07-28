"""
docx 寫回器 (Format-preserving Writer)
------------------------------------------------
最大的坑：Word 會把一句話拆成好幾個 <w:r> run（拼字檢查、修訂 ID 都會造成切割），
所以「你看得到的字串」在 XML 裡往往不是連續的。直接用 cell.text = "..." 會把
儲存格內所有格式（字型、大小、置中）一次清光，出來的履歷會很醜。

這裡的作法：
  1. 先把同格式的相鄰 run 合併（coalesce），讓字串變成可搜尋
  2. 只改動目標 run 的 <w:t> 文字，其餘 XML 一律不碰
  3. 空白格沒有任何 run 可繼承時，從同一列借一個 run 的 <w:rPr> 複製過來

--highlight 模式會把填入的字加上黃色底色，方便人工快速掃過一遍再取消標示。
"""

from __future__ import annotations

import copy
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table

from .document import (BLANK_RUN_RE, CHECKBOX_CHARS, CHECKED_CHARS,
                       _grid, cell_text, iter_block_items)

CHECK_MAP = {"□": "■", "☐": "☑", "▢": "■", "◻": "◼", "○": "●", "〇": "●", "◯": "●"}


# --------------------------------------------------------------------------
# run 層級的工具
# --------------------------------------------------------------------------
def coalesce_runs(paragraph) -> None:
    """合併相鄰且格式相同的 run，讓文字變成連續可搜尋（不改變外觀）。"""
    runs = paragraph.runs
    i = 0
    while i < len(runs) - 1:
        a, b = runs[i], runs[i + 1]
        rpr_a = a._element.find(qn("w:rPr"))
        rpr_b = b._element.find(qn("w:rPr"))
        same = (rpr_a is None and rpr_b is None) or (
            rpr_a is not None and rpr_b is not None
            and rpr_a.xml == rpr_b.xml)
        has_special = any(b._element.find(qn(t)) is not None
                          for t in ("w:drawing", "w:fldChar", "w:instrText", "w:br", "w:tab"))
        if same and not has_special:
            a.text = a.text + b.text
            b._element.getparent().remove(b._element)
            runs = paragraph.runs
        else:
            i += 1


def _set_run_text(run, text: str, highlight: bool = False) -> None:
    run.text = text
    if highlight:
        rpr = run._element.get_or_add_rPr()
        for old in rpr.findall(qn("w:highlight")):
            rpr.remove(old)
        el = rpr.makeelement(qn("w:highlight"), {qn("w:val"): "yellow"})
        rpr.append(el)


def _clone_rpr(src_run, dst_run) -> None:
    rpr = src_run._element.find(qn("w:rPr"))
    if rpr is not None:
        dst_run._element.insert(0, copy.deepcopy(rpr))


def _donor_run(table: Table, row_idx: int):
    """從同一列找一個有文字的 run，借它的字型設定。"""
    try:
        cells = table.rows[row_idx].cells
    except IndexError:
        return None
    for cell in cells:
        for p in cell.paragraphs:
            for r in p.runs:
                if r.text.strip():
                    return r
    return None


def _write_into_cell(table: Table, row: int, col: int, text: str,
                     highlight: bool) -> bool:
    grid = _grid(table)
    if row >= len(grid) or col >= len(grid[row]) or grid[row][col] is None:
        return False
    cell = grid[row][col]
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    coalesce_runs(para)
    if para.runs:
        _set_run_text(para.runs[0], text, highlight)
        for r in para.runs[1:]:
            r.text = ""
    else:
        run = para.add_run()
        donor = _donor_run(table, row)
        if donor is not None:
            _clone_rpr(donor, run)
        _set_run_text(run, text, highlight)
    return True


def _replace_span(para, start: int, end: int, text: str, highlight: bool) -> bool:
    """把段落文字的 [start, end) 區間換成 text，只動到牽涉到的 run。"""
    coalesce_runs(para)
    total = len(para.text)
    if start >= total:                      # 純追加：「可到職日：」後面直接接上
        new = para.add_run()                # 另開新 run，避免把標籤一起上色
        if para.runs and len(para.runs) > 1:
            _clone_rpr(para.runs[-2], new)
        _set_run_text(new, text, highlight)
        return True

    pos = 0
    written = False
    for run in list(para.runs):
        rlen = len(run.text)
        r_start, r_end = pos, pos + rlen
        pos = r_end
        if r_end <= start or r_start >= end:
            continue
        cut_a = max(0, start - r_start)
        cut_b = min(rlen, end - r_start)
        head, tail = run.text[:cut_a], run.text[cut_b:]
        if not written:
            _set_run_text(run, head + text + tail, highlight)
            written = True
        else:
            run.text = head + tail
    return written


def _fill_inline(para, text: str, highlight: bool, blank_index: int = 0,
                 span: Optional[tuple] = None) -> bool:
    """處理『姓名：______』『… 關係：___ 電話：___』或『可到職日：』結尾。"""
    # span 是既有值的字元範圍（來自 include_filled 抽取）。沒有它的話，
    # 一個已經填好的段落找不到空白也沒有結尾冒號，只會把新值附加在後面。
    if span is not None:
        return _replace_span(para, span[0], span[1], text, highlight)

    full = para.text
    blanks = list(BLANK_RUN_RE.finditer(full))
    if blanks:
        if blank_index >= len(blanks):
            return False
        m = blanks[blank_index]
        return _replace_span(para, m.start(), m.end(), text, highlight)
    m = re.search(r"[:：][ \u3000]*$", full)
    if m:
        return _replace_span(para, m.end(), len(full), text, highlight)
    return _replace_span(para, len(full), len(full), text, highlight)


def _fill_checkbox(para, option: str, highlight: bool) -> bool:
    full = para.text
    idx = full.find(option)
    if idx < 0:
        return False
    for j in range(idx - 1, max(-1, idx - 4), -1):
        ch = full[j]
        if ch in CHECKBOX_CHARS:
            return _replace_span(para, j, j + 1, CHECK_MAP.get(ch, "■"), highlight)
    return False


# --------------------------------------------------------------------------
# 內容控制項 / 舊式表單欄位
# --------------------------------------------------------------------------
def _fill_sdt(doc, index: int, text: str) -> bool:
    sdts = list(doc.element.body.iter(qn("w:sdt")))
    if index >= len(sdts):
        return False
    sdt = sdts[index]
    pr = sdt.find(qn("w:sdtPr"))
    if pr is not None:                       # 清掉「按此輸入」的預留位置狀態
        for ph in pr.findall(qn("w:showingPlcHdr")):
            pr.remove(ph)
    content = sdt.find(qn("w:sdtContent"))
    if content is None:
        return False
    tnodes = list(content.iter(qn("w:t")))
    if tnodes:
        tnodes[0].text = text
        for t in tnodes[1:]:
            t.text = ""
        return True
    return False


def _fill_formfield(doc, index: int, text: str) -> bool:
    ffs = list(doc.element.body.iter(qn("w:ffData")))
    if index >= len(ffs):
        return False
    run = ffs[index].getparent().getparent()          # ffData -> fldChar -> r
    parent = run.getparent()
    children = list(parent)
    try:
        i = children.index(run)
    except ValueError:
        return False
    state = "begin"
    for node in children[i + 1:]:
        fld = node.find(qn("w:fldChar")) if node.tag == qn("w:r") else None
        if fld is not None:
            ftype = fld.get(qn("w:fldCharType"))
            if ftype == "separate":
                state = "result"
                continue
            if ftype == "end":
                break
        if state == "result" and node.tag == qn("w:r"):
            tnodes = list(node.iter(qn("w:t")))
            if tnodes:
                tnodes[0].text = text
                for t in tnodes[1:]:
                    t.text = ""
                return True
    return False


# --------------------------------------------------------------------------
# 對外主函式
# --------------------------------------------------------------------------
def apply_ops(src_path: str, out_path: str, ops: List[Any],
              highlight: bool = False) -> Dict[str, Any]:
    doc = Document(src_path)
    ok, fail = [], []

    # 同一段落有多個填空時，必須由後往前寫，
    # 否則填完第 1 個空格後，後面空格的字元位置就跑掉了。
    ops = sorted(ops, key=lambda o: -o.slot.loc.get("blank_index", 0))

    for op in ops:
        slot = op.slot
        loc = slot.loc
        kind = slot.kind
        done = False
        try:
            if kind == "sdt":
                done = _fill_sdt(doc, loc["sdt_index"], op.value)
            elif kind == "formfield":
                done = _fill_formfield(doc, loc["ff_index"], op.value)
            elif kind == "cell":
                done = _write_into_cell(doc.tables[loc["table"]],
                                        loc["row"], loc["col"], op.value, highlight)
            elif kind == "inline":
                bi = loc.get("blank_index", 0)
                span = ((loc["value_start"], loc["value_end"])
                        if "value_start" in loc else None)
                if loc.get("in_cell"):
                    grid = _grid(doc.tables[loc["table"]])
                    cell = grid[loc["row"]][loc["col"]]
                    done = _fill_inline(cell.paragraphs[-1], op.value, highlight, bi, span)
                else:
                    done = _fill_inline(doc.paragraphs[loc["para"]], op.value,
                                        highlight, bi, span)
            elif kind == "checkbox":
                if "para" in loc:
                    done = _fill_checkbox(doc.paragraphs[loc["para"]], op.value, highlight)
                else:
                    grid = _grid(doc.tables[loc["table"]])
                    cell = grid[loc["row"]][loc["col"]]
                    for p in cell.paragraphs:
                        if _fill_checkbox(p, op.value, highlight):
                            done = True
                            break
        except Exception as e:                    # 單一格失敗不影響其他欄位
            fail.append({"slot": slot.id, "error": str(e)})
            continue
        (ok if done else fail).append(
            {"slot": slot.id, "field": op.field_key}
            if done else {"slot": slot.id, "error": "定位失敗"})

    doc.save(out_path)
    return {"written": len(ok), "failed": len(fail), "ok": ok, "fail": fail}
