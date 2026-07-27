"""產生一份仿台灣公司格式的 Word 履歷表，用來測試整條流程。"""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build(path="sample_resume_form.docx"):
    doc = Document()
    doc.styles["Normal"].font.name = "DFKai-SB"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_paragraph("○○科技股份有限公司　應徵人員資料表")
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(16)

    # --- 基本資料表：標籤在左、空白在右 ---
    t = doc.add_table(rows=6, cols=4)
    t.style = "Table Grid"
    layout = [
        ["姓　名", "", "性　別", "□男　□女"],
        ["出生年月日", "", "行動電話", ""],
        ["電子信箱", "", "兵役狀況", "□役畢　□免役　□未役"],
        ["戶籍地址", "", "", ""],
        ["通訊地址", "", "", ""],
        ["應徵職務", "", "希望待遇", ""],
    ]
    for r, row in enumerate(layout):
        for c, val in enumerate(row):
            t.cell(r, c).text = val
    t.cell(3, 1).merge(t.cell(3, 3))
    t.cell(4, 1).merge(t.cell(4, 3))

    doc.add_paragraph()

    # --- 學歷表：標籤在表頭列、資料在下面幾列 ---
    doc.add_paragraph("一、學歷").runs[0].bold = True
    edu = doc.add_table(rows=4, cols=4)
    edu.style = "Table Grid"
    for c, head in enumerate(["學校名稱", "科系", "學位", "就學期間"]):
        edu.cell(0, c).text = head

    doc.add_paragraph()

    # --- 經歷表 ---
    doc.add_paragraph("二、工作經歷").runs[0].bold = True
    exp = doc.add_table(rows=4, cols=4)
    exp.style = "Table Grid"
    for c, head in enumerate(["公司名稱", "職稱", "任職期間", "離職原因"]):
        exp.cell(0, c).text = head

    doc.add_paragraph()

    # --- 段落填空 ---
    doc.add_paragraph("三、其他").runs[0].bold = True
    doc.add_paragraph("語文能力：____________________________")
    doc.add_paragraph("專業證照：____________________________")
    doc.add_paragraph("緊急聯絡人：____________  關係：________  緊急聯絡電話：____________")
    doc.add_paragraph("可到職日：")

    doc.add_paragraph()
    doc.add_paragraph("四、自傳").runs[0].bold = True
    bio = doc.add_table(rows=1, cols=1)
    bio.style = "Table Grid"
    bio.cell(0, 0).text = ""

    doc.add_paragraph()
    doc.add_paragraph("※以下由人事單位填寫")
    hr = doc.add_table(rows=1, cols=2)
    hr.style = "Table Grid"
    hr.cell(0, 0).text = "面試評語"
    hr.cell(0, 1).text = ""

    doc.save(path)
    print(f"✓ 已產生 {path}")


if __name__ == "__main__":
    build(sys.argv[1] if len(sys.argv) > 1 else "sample_resume_form.docx")
